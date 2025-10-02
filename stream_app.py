
import streamlit as st
import textwrap
import os
import io
import contextlib
import fitz
import json
import shutil
import gc
import time
from docx import Document as DocxDocument
from datetime import datetime
from collections import OrderedDict

import os

# Not needed if ENV set in Dockerfile, but safe fallback
os.environ.setdefault("STREAMLIT_CONFIG_DIR", "/tmp/.streamlit")
os.environ.setdefault("TRANSFORMERS_CACHE", "/tmp/hf_cache")
os.environ.setdefault("HF_HOME", "/tmp/hf_cache")

os.makedirs(os.environ["STREAMLIT_CONFIG_DIR"], exist_ok=True)
os.makedirs(os.environ["TRANSFORMERS_CACHE"], exist_ok=True)
os.makedirs(os.environ["HF_HOME"], exist_ok=True)

# --- LangChain and AI Imports ---
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnableSequence
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.agents import initialize_agent, Tool
from langchain.agents.agent_types import AgentType
from langchain_community.tools.ddg_search import DuckDuckGoSearchRun
from langchain_community.vectorstores import Chroma
from langchain.docstore.document import Document
# from chromadb import PersistentClient
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from asteval import Interpreter
import psycopg2
from langchain_community.vectorstores.pgvector import PGVector
from langchain_community.vectorstores import FAISS

# =============================================================================
# --- ⚙️ INTEGRATED BACKEND LOGIC ---
# =============================================================================

# --- 1. Configuration ---
# SECURELY CONFIGURE THE LLM AND EMBEDDINGS
try:
    llm = ChatOpenAI(
        model="provider-3/gpt-5-nano",
        temperature=0.3,
        streaming=True,
        base_url="https://api.a4f.co/v1",
        api_key="ddc-a4f-813de865fbe3421fba1419c81047de00"
    )
    embedding_model_name = "sentence-transformers/all-MiniLM-L6-v2"
    embeddings = HuggingFaceEmbeddings(
        model_name=embedding_model_name,
        model_kwargs={'device': 'cpu'})
except KeyError:
    st.error("Missing OPENAI_API_KEY. Please add it to your Hugging Face Space Secrets.")
    st.stop()
except Exception as e:
    st.error(f"Error initializing AI models: {e}")
    st.stop()


# --- 2. Goal Decomposer ---
goal_prompt = PromptTemplate(
    input_variables=["goal", "industry", "urgency"],
    template="""
You are a top-tier enterprise AI strategist with deep understanding of business operations across domains.
Your task is to take the following high-level goal and generate 5–10 detailed, logically ordered, and actionable steps to achieve it.
📌 Context:
- Industry: {industry}
  (If 'Other' or unspecified, analyze the goal and intelligently infer the most relevant industry context.)
- Urgency Level: {urgency}
- Goal: "{goal}"
✅ Format Instructions:
- Numbered list (1., 2., 3., ...)
- Each task should begin with a verb (e.g., Research, Analyze, Design, Implement)
- Keep tasks concise but specific
- No additional explanations or extra text
Respond only with the list of tasks.
"""
)
goal_chain = goal_prompt | llm

def decompose_goal(goal: str, industry: str = "General", urgency: str = "Medium") -> list:
    try:
        response = goal_chain.invoke({"goal": goal, "industry": industry, "urgency": urgency})
        output = response.content if hasattr(response, "content") else str(response)
        tasks = [line.lstrip("0123456789.- ").strip() for line in output.split("\n") if line.strip()]
        return tasks
    except Exception as e:
        st.error(f"Error in decompose_goal: {e}")
        return []


# --- 3. Task Executor ---
search = DuckDuckGoSearchRun()

def run_python_code(code: str) -> str:
    output = io.StringIO()
    try:
        with contextlib.redirect_stdout(output):
            exec(code, {})
        return output.getvalue().strip() or "Code executed successfully."
    except Exception as e:
        return f"Error: {str(e)}"

tools = [
    Tool(
        name="Web Search",
        func=search.run,
        description="Use for answering questions about current events or web data"
    ),
    Tool(
        name="Python Calculator",
        func=run_python_code,
        description="Useful for math, code, or data calculations"
    )
]
router_agent = initialize_agent(
    tools=tools, llm=llm, agent=AgentType.OPENAI_FUNCTIONS, verbose=True, handle_parsing_errors=True
)

def execute_task(task: str, goal: str, industry: str = "General", urgency: str = "Medium") -> dict:
    try:
        prompt = f"""
You are a specialized Enterprise AI Agent designed to autonomously execute high-impact tasks in business environments.
🎯 Goal: "{goal}"
📌 Task: "{task}"
🏭 Industry: {industry}
⏱️ Urgency: {urgency}
Instructions:
- Think like a domain expert in the given industry.
- Adjust the depth and immediacy of the solution based on the urgency level.
- Use tools (search, calculator) if required to complete the task.
- Be clear, factual, and directly useful. Avoid unnecessary commentary.
Return only the result of the task execution.
"""
        result = router_agent.invoke({"input": prompt.strip()})
        return {"task": task, "result": result, "status": "success"}
    except Exception as e:
        return {"task": task, "result": str(e), "status": "failed"}

# --- 4. Task Planner ---
priority_prompt = PromptTemplate(
    input_variables=["tasks"],
    template="""
You are a senior enterprise operations strategist AI.
You are given a list of operational tasks for a business objective:
{tasks}
Your job is to rank these tasks by execution priority, based on the following criteria:
- Logical dependencies: Which tasks depend on others?
- Business impact: Which tasks unlock the most value or mitigate the highest risk?
- Parallelizability: Which tasks can be run in parallel vs. require sequence?
- Time-sensitivity: Which tasks are urgent or unblock other workflows?
Return ONLY a JSON list in the following format:
[
  {{"task": "task name", "priority": 1}},
  ...
]
Note:
- 1 = highest priority, higher numbers = lower priority
- Do not explain. Only output clean JSON list.
"""
)
priority_chain = priority_prompt | llm

def prioritize_tasks(task_list: list) -> list:
    tasks_str = "\n".join(f"- {task}" for task in task_list)
    try:
        response = priority_chain.invoke({"tasks": tasks_str})
        task_data = json.loads(response.content if hasattr(response, "content") else response)
        return task_data
    except Exception as e:
        st.error(f"❌ Error parsing priority response: {e}")
        return [{"task": task, "priority": 99} for task in task_list]


import psycopg2
from langchain_community.vectorstores.pgvector import PGVector

# --- 5. Memory Manager (Updated for Neon DB) ---
def get_memory_vectorstore():
    """Connects to the Neon database using a secure connection string from Streamlit secrets."""
    try:
        connection_string = st.secrets["NEON_DB_URL"]
    except KeyError:
        st.error("Missing NEON_DB_URL in Streamlit secrets. Please add it to your app's secrets.")
        st.stop()
    
    return PGVector(
        connection_string=connection_string,
        embedding_function=embeddings,
        collection_name="memory"
    )

def save_to_memory(goal: str, task: str, result: str,
                   evaluation: str, industry: str, urgency: str, status: str = "completed"):
    iso_timestamp = datetime.utcnow().isoformat()
    readable_timestamp = datetime.now().strftime("%d %b %Y, %I:%M %p")

    metadata = {
        "goal": goal,
        "task": task,
        "industry": industry,
        "urgency": urgency,
        "result": result,
        "evaluation": evaluation,
        "status": status,
        "timestamp": readable_timestamp,
        "iso_timestamp": iso_timestamp,
    }

    doc = Document(page_content=result, metadata=metadata)
    vectorstore = get_memory_vectorstore()

    similar_docs = vectorstore.similarity_search(task, k=1)
    if similar_docs and similar_docs[0].page_content == doc.page_content:
        st.warning("⚠️ Task already exists in memory. Skipping save.")
        return

    vectorstore.add_documents([doc])
    st.info("✅ Execution memory saved successfully.")

def clear_memory():
    """Deletes the 'memory' table from the database."""
    try:
        conn = psycopg2.connect(st.secrets["NEON_DB_URL"])
        conn.set_isolation_level(psycopg2.extensions.ISOLATION_LEVEL_AUTOCOMMIT)
        cursor = conn.cursor()
        cursor.execute("DROP TABLE IF EXISTS langchain_pg_embedding;")
        conn.close()
        # st.success("✅ Memory cleared successfully.")
    except Exception as e:
        st.error(f"⚠️ Error while clearing memory: {e}")

def get_memory_entries() -> list:
    try:
        vectorstore = get_memory_vectorstore()
        docs = vectorstore.similarity_search(" ", k=100)
        memory_data = []
        for doc in docs:
            metadata = doc.metadata
            memory_data.append({
                "goal": metadata.get("goal", ""),
                "task": metadata.get("task", doc.page_content),
                "industry": metadata.get("industry", "General"),
                "urgency": metadata.get("urgency", "Normal"),
                "result": metadata.get("result", doc.page_content),
                "evaluation": metadata.get("evaluation", ""),
                "status": metadata.get("status", "completed"),
                "timestamp": metadata.get("timestamp", "")
            })
        return memory_data
    except Exception as e:
        return [f"⚠️ Error reading memory: {str(e)}"]
    
# =============================================================================
# --- STREAMLIT UI CODE ---
# =============================================================================

# Global State Management
if 'session_text' not in st.session_state:
    st.session_state.session_text = None
if "planned_tasks" not in st.session_state:
    st.session_state["planned_tasks"] = []
if "memory_data" not in st.session_state:
    st.session_state.memory_data = []

st.set_page_config(page_title="EAOA Agent", page_icon="🤖", layout="wide")

st.markdown(
    """
    <style>
    .reportview-container { background: #f0f2f6; }
    .sidebar .sidebar-content { background-color: #ffffff; padding-top: 2rem; }
    h1, h2, h3, h4, h5, h6 { color: #1a1a2e; font-family: 'Segoe UI', sans-serif; }
    .st-emotion-cache-13ln4tx { padding: 2rem; border-radius: 12px; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1); background-color: #ffffff; border: 1px solid #e0e0e0; margin-bottom: 20px; }
    .st-emotion-cache-9r12zj { border-radius: 8px; border: 1px solid #4a90e2; background-color: #4a90e2; color: white; transition: all 0.2s ease-in-out; font-weight: 600; padding: 10px 20px; }
    .st-emotion-cache-9r12zj:hover { background-color: #3b76c2; border-color: #3b76c2; transform: translateY(-2px); }
    .st-emotion-cache-9r12zj:active { background-color: #2c5a9b; border-color: #2c5a9b; transform: translateY(0); }
    .st-emotion-cache-j7qwjs { border-radius: 8px; border: 1px solid #dcdcdc; }
    .stAlert { border-radius: 8px; }
    .stAlert.success { background-color: #d4edda; color: #155724; }
    .stAlert.error { background-color: #f8d7da; color: #721c24; }
    .stAlert.info { background-color: #d1ecf1; color: #0c5460; }
    .st-emotion-cache-14n935g > div > div > input { font-family: 'Courier New', monospace; font-size: 2.5rem !important; font-weight: bold; color: #1a1a2e; padding: 20px; text-align: right; border: 2px solid #4a90e2; border-radius: 8px; min-height: 80px; }
    .st-emotion-cache-11r009h { border-radius: 8px; border: 1px solid #e0e0e0; background-color: #ffffff; }
    </style>
    """,
    unsafe_allow_html=True
)

st.sidebar.title("Configuration")
st.sidebar.markdown("---")
st.sidebar.info("EAOA Agent v1.0")
st.sidebar.markdown("Developed for Resume Showcase")
st.sidebar.text_input("API Endpoint", value="Local Integrated", disabled=True)

st.title("Enterprise Autonomous Operations Agent")
st.subheader("An intelligent agent for goal-oriented task execution")

tabs = st.tabs(["🎯 Goal Executor", "📝 Document Q&A", "🧮 Calculator", "🧠 Memory"])

# -----------------------------------------
# Tool 1: Goal Executor
# -----------------------------------------
with tabs[0]:
    st.markdown("### Plan & Execute Tasks")
    st.info("First, define a high-level business goal. The agent will then generate a list of actionable tasks. Finally, you can execute these tasks one by one.")
    
    with st.container():
        st.subheader("Define a New Goal")
        with st.form("goal_form"):
            goal = st.text_area("High-level Goal", placeholder="e.g., 'Develop a marketing strategy for a new product launch.'")
            col1, col2 = st.columns(2)
            with col1:
                industry = st.selectbox("Industry", ["General", "Education", "Healthcare", "Finance", "Retail", "Tech", "Manufacturing", "Agriculture", "Energy", "Logistics", "Other"])
            with col2:
                urgency = st.selectbox("Urgency", ["Low", "Medium", "High", "Critical"])
            
            submitted = st.form_submit_button("Plan Tasks", use_container_width=True)

        if submitted:
            with st.spinner("Planning tasks based on your goal..."):
                try:
                    tasks = decompose_goal(goal, industry, urgency)
                    prioritized_tasks = prioritize_tasks(tasks)
                    
                    for task in prioritized_tasks:
                        task["industry"] = industry
                        task["urgency"] = urgency
                    
                    st.session_state["planned_tasks"] = [{"task " + str(i+1): task['task'], "priority": task['priority']} for i, task in enumerate(prioritized_tasks)]
                    st.session_state["last_goal"] = goal
                    st.session_state["industry"] = industry
                    st.session_state["urgency"] = urgency
                    st.session_state["result"] = "Tasks planned successfully."
                    st.success("✅ Tasks planned successfully!")

                except Exception as e:
                    st.error(f"Error planning tasks: {e}")

    if "planned_tasks" in st.session_state and st.session_state["planned_tasks"]:
        with st.expander("📋 View Planned Tasks Summary", expanded=True):
            st.markdown(f"**Goal:** `{st.session_state.get('last_goal', 'N/A')}`")
            st.markdown(f"**Industry:** `{st.session_state.get('industry', 'N/A')}` | **Urgency:** `{st.session_state.get('urgency', 'N/A')}`")
            st.markdown("---")
            st.markdown("#### Generated Tasks")
            for i, task in enumerate(st.session_state["planned_tasks"], start=1):
                task_key = f"task {i}"
                task_text = task.get(task_key, "Unknown Task")
                priority = task.get("priority", "N/A")
                st.markdown(f"**Task {i}:**")
                st.markdown(task_text)
                st.markdown(f"**Priority:** {priority}")
                st.markdown("---")
            
            if st.session_state.get("result"):
                st.markdown("---")
                st.markdown("#### Final Result")
                st.markdown(st.session_state["result"])

    st.markdown("---")
    
    st.subheader("Execute a Specific Task")
    with st.container():
        task_number = st.number_input("Enter the task number to execute", min_value=1, step=1, key="task_num_exec")
        if st.button("Run Task", use_container_width=True):
            with st.spinner(f"Executing Task {task_number}..."):
                try:
                    tasks = st.session_state.get("planned_tasks")
                    goal = st.session_state.get("last_goal")
                    if not tasks or not goal:
                        st.error("❗ No tasks found. Please submit a goal first.")
                    elif task_number > len(tasks):
                        st.error(f"❌ Invalid task number. Only {len(tasks)} tasks available.")
                    else:
                        task_data = tasks[task_number - 1]
                        task_text = task_data[f"task {task_number}"]
                        
                        execution_result = execute_task(
                            task=task_text,
                            goal=goal,
                            industry=st.session_state.get("industry", "General"),
                            urgency=st.session_state.get("urgency", "Medium")
                        )

                        raw_result = execution_result.get("result", {})
                        clean_output = raw_result.get("output", "No output returned.")
                        
                        save_to_memory(
                            goal=goal,
                            task=task_text,
                            result=clean_output,
                            evaluation="Executed from Streamlit UI",
                            industry=st.session_state.get("industry", "General"),
                            urgency=st.session_state.get("urgency", "Medium")
                        )
                        st.session_state["execution_result"] = {
                            "goal": goal,
                            f"task {task_number}": task_text,
                            "result": clean_output
                        }
                        st.session_state["last_executed_task"] = task_number

                except Exception as e:
                    st.error(f"Error executing task: {e}")
                    st.session_state["execution_result"] = None

    if st.session_state.get("execution_result"):
        result = st.session_state["execution_result"]
        task_number = st.session_state["last_executed_task"]
        goal = result.get("goal", "N/A")
        task_key = f"task {task_number}"
        task_text = result.get(task_key, "N/A")
        output = result.get("result", "No result returned.")
        cleaned_output = textwrap.dedent(output).strip()
        
        with st.container(border=True):
            st.success(f"✅ Task {task_number} Executed Successfully!")
            st.markdown(f"**Goal:** `{goal}`")
            st.markdown(f"**Task:** `{task_text}`")
            st.markdown("---")
            st.markdown("#### Execution Result:")
            st.markdown(cleaned_output)

# -----------------------------------------
# Tool 2: Document Q&A
# -----------------------------------------
with tabs[1]:
    st.markdown("### Document Question & Answer")
    st.info("Upload a document, then ask questions to get quick answers without saving the content to memory.")
    
    with st.container():
        st.subheader("Upload a Document")
        uploaded_file = st.file_uploader("Upload a .txt, .pdf, or .docx file", type=["txt", "pdf", "docx"])
        if uploaded_file:
            if st.button("Upload Document", use_container_width=True):
                with st.spinner("Uploading document..."):
                    try:
                        filename = uploaded_file.name
                        ext = os.path.splitext(filename)[-1].lower()
                        
                        if ext == ".txt":
                            st.session_state.session_text = uploaded_file.getvalue().decode("utf-8")
                        elif ext == ".pdf":
                            doc = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
                            st.session_state.session_text = "".join([page.get_text() for page in doc])
                            doc.close()
                        elif ext == ".docx":
                            doc = DocxDocument(uploaded_file)
                            st.session_state.session_text = "\n".join([para.text for para in doc.paragraphs])
                        else:
                            st.error("❌ Unsupported file type.")
                            st.session_state.session_text = None
                        
                        if st.session_state.session_text:
                            st.success(f"✅ Document '{filename}' uploaded for session. You can now ask questions.")
                    
                    except Exception as e:
                        st.error(f"❌ Upload failed: {str(e)}")
                        st.session_state.session_text = None

    st.markdown("---")
    
    with st.container():
        st.subheader("Ask a Question")
        query = st.text_input("Your question about the document", placeholder="e.g., 'What is the main topic of the document?'")
        
        if st.button("Get Answer", use_container_width=True):
            if query and st.session_state.session_text:
                with st.spinner("Analyzing document and generating answer..."):
                    try:
                        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                        docs = text_splitter.create_documents([st.session_state.session_text])

                        # ✅ FIX: Use in-memory Chroma
                        vectordb = FAISS.from_documents(docs, embedding=embeddings)


                        retriever = vectordb.as_retriever()
                        qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)
                        result = qa.run(query)

                        st.success("Answer:")
                        st.write(result)
                    except Exception as e:
                        st.error(f"❌ Query failed: {str(e)}")
            elif not st.session_state.session_text:
                st.warning("⚠️ No document uploaded. Please upload a document first.")
            else:
                st.warning("Please enter a question.")

# -----------------------------------------
# Tool 3: Advanced Calculator
# -----------------------------------------
with tabs[2]:
    st.markdown("### Advanced Calculator")
    st.info("Enter your mathematical expression directly below. The calculator supports functions like `pow()`, `sqrt()`, and more.")

    with st.container():
        st.subheader("Evaluate Expression")
        expression = st.text_input(
            "Enter expression",
            placeholder="e.g., (15 * 5) + pow(2, 4)",
            label_visibility="collapsed"
        )
        
        if st.button("Calculate", use_container_width=True):
            if expression:
                with st.spinner("Calculating..."):
                    try:
                        ae = Interpreter()
                        result = ae(expression)
                        if ae.error:
                            raise Exception(ae.error[0].get_error())
                        st.success(f"Result: {result}")
                    except Exception as e:
                        st.error(f"Calculation failed: {e}")
            else:
                st.warning("Please enter an expression.")
        
# -----------------------------------------
# Tool 4: Memory Viewer
# -----------------------------------------
with tabs[3]:
    st.markdown("### Memory Viewer & Management")
    st.info("View a log of all executed tasks and their results. This memory is persistent until you clear it.")
    
    with st.container():
        st.subheader("Manage Memory")
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("🔄 Refresh Memory", use_container_width=True):
                with st.spinner("Fetching memory..."):
                    try:
                        st.session_state.memory_data = get_memory_entries()
                        st.success("Memory refreshed.")
                    except Exception as e:
                        st.error(f"Failed to fetch memory: {e}")
            
        with col2:
            if st.button("🗑️ Clear Memory", use_container_width=True):
                with st.spinner("Clearing memory..."):
                    try:
                        clear_memory()
                        st.session_state.memory_data = []
                        st.success("✅ Memory cleared successfully.")
                    except Exception as e:
                        st.error(f"Memory clear failed: {e}")

    memory = st.session_state.get("memory_data", [])

    # This is the corrected 'if' statement
    if memory and isinstance(memory[0], dict):
        st.markdown("---")
        st.subheader("Stored Task History")
        for idx, entry in enumerate(memory, start=1):
            with st.container():
                raw_ts = entry.get('timestamp', 'N/A')
                st.markdown(
                    f"#### Memory Entry #{idx} <span style='float:right;color:gray;font-size:14px'>🗓️ {raw_ts}</span>",
                    unsafe_allow_html=True
                )
                st.markdown(f"**Goal:** `{entry.get('goal', 'N/A')}`")
                
                task_text = entry.get('task', 'N/A')
                st.markdown(f"**Task:** `{task_text}`")
                
                st.markdown(f"**Result:**")
                st.markdown(entry.get('result', 'N/A'))
                st.markdown("---")
    else:
        st.info("No memory entries found. Click 'Refresh Memory' to view any saved tasks.")